from __future__ import annotations

import struct
import zlib
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
FIXTURES = ROOT / "test-fixtures"


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(10)

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    bullet.paragraph_format.line_spacing = 1.25


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def create_docx() -> None:
    document = Document()
    configure_docx_styles(document)
    document.add_heading("Knowledge Base Guide", level=1)
    paragraph = document.add_paragraph("Read the ")
    add_hyperlink(paragraph, "operations handbook", "https://example.com/handbook")
    paragraph.add_run(" before publishing a document.")
    document.add_paragraph("Upload a source file", style="List Bullet")
    document.add_paragraph("Review normalized Markdown", style="List Bullet")

    table = document.add_table(rows=3, cols=2)
    table.autofit = False
    widths = (Inches(2), Inches(4.5))
    values = (("Stage", "Result"), ("Parse", "Markdown"), ("Index", "Searchable"))
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.width = widths[column_index]
            cell.text = values[row_index][column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    document.add_picture(BytesIO(solid_png(80, 60)), width=Inches(0.8))

    document.save(FIXTURES / "parser-sample.docx")


def png_chunk(kind: bytes, payload: bytes) -> bytes:
    return struct.pack(">I", len(payload)) + kind + payload + struct.pack(">I", zlib.crc32(kind + payload))


def solid_png(width: int, height: int) -> bytes:
    rows = b"".join(b"\x00" + bytes((35, 130, 190)) * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + png_chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + png_chunk(b"IDAT", zlib.compress(rows))
        + png_chunk(b"IEND", b"")
    )


def create_pdf() -> None:
    output = FIXTURES / "parser-sample.pdf"
    document = canvas.Canvas(str(output), pagesize=letter)
    document.setTitle("Knowledge Base Parser Sample")
    document.setFont("Helvetica-Bold", 18)
    document.drawString(72, 720, "Knowledge Base PDF")
    document.setFont("Helvetica", 11)
    document.drawString(72, 690, "This is the first page used for parser verification.")
    document.drawImage(ImageReader(BytesIO(solid_png(120, 80))), 72, 560, width=120, height=80)
    document.showPage()
    document.setFont("Helvetica-Bold", 18)
    document.drawString(72, 720, "Operations Checklist")
    document.setFont("Helvetica", 11)
    document.drawString(72, 690, "Confirm that page anchors preserve the source page number.")
    document.save()


def main() -> None:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    create_docx()
    create_pdf()
    print(f"Generated parser fixtures in {FIXTURES}")


if __name__ == "__main__":
    main()
